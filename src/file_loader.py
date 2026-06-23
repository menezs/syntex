from pathlib import Path
from typing import List
from src.models import Document


class FileLoader:
    def __init__(self, answers_dir: str):
        self.answers_dir = Path(answers_dir)

    def load(self) -> List[Document]:
        documents = []

        if not self.answers_dir.exists():
            raise FileNotFoundError(f"Directory not found: {self.answers_dir}")

        for filepath in sorted(self.answers_dir.iterdir()):
            if filepath.is_dir():
                continue

            suffix = filepath.suffix.lower()
            if suffix not in (".pdf", ".docx", ".md"):
                continue

            try:
                content = self._load_file(filepath, suffix)
                doc = Document(
                    path=str(filepath),
                    filename=filepath.name,
                    content=content,
                    metadata={"path": str(filepath), "format": suffix}
                )
                documents.append(doc)
                print(f"  Loaded: {filepath.name} ({len(content)} chars)")
            except Exception as e:
                print(f"  Error loading {filepath.name}: {e}")

        return documents

    def _load_file(self, filepath: Path, suffix: str) -> str:
        if suffix == ".pdf":
            return self._load_pdf(filepath)
        elif suffix == ".docx":
            return self._load_docx(filepath)
        elif suffix == ".md":
            return self._load_markdown(filepath)
        else:
            raise ValueError(f"Unsupported format: {suffix}")

    def _load_pdf(self, filepath: Path) -> str:
        import pymupdf4llm
        md_text = pymupdf4llm.to_markdown(str(filepath))
        return md_text

    def _load_docx(self, filepath: Path) -> str:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        doc = DocxDocument(str(filepath))
        paragraphs = []

        for paragraph in doc.paragraphs:
            text = self._extract_docx_paragraph(paragraph)
            if text.strip():
                paragraphs.append(text)

        return "\n\n".join(paragraphs)

    def _extract_docx_paragraph(self, paragraph) -> str:
        from docx.oxml.ns import qn
        parts = []

        for child in paragraph._element:
            if child.tag == qn("w:r"):
                parts.append(self._extract_run_text(child))
            elif child.tag == qn("w:hyperlink"):
                for run_elem in child.findall(qn("w:r")):
                    parts.append(self._extract_run_text(run_elem))

        return "".join(parts)

    def _extract_run_text(self, run_elem) -> str:
        from docx.oxml.ns import qn
        from docx.text.run import Run

        t_elem = run_elem.find(qn("w:t"))
        if t_elem is None:
            return ""

        text = t_elem.text or ""

        rpr = run_elem.find(qn("w:rPr"))
        if rpr is not None:
            vert_align = rpr.find(qn("w:vertAlign"))
            if vert_align is not None:
                val = vert_align.get(qn("w:val"))
                if val == "superscript":
                    return f"<sup>{text}</sup>"
                elif val == "subscript":
                    return f"<sub>{text}</sub>"

        return text

    def _load_markdown(self, filepath: Path) -> str:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
